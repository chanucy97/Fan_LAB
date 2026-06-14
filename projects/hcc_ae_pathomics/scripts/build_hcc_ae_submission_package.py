from __future__ import annotations

import argparse
import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_DIR = Path(__file__).resolve().parents[1]
CONTROLLED_ROOT = Path(os.environ.get("HCC_AE_ROOT", Path.cwd())).resolve()
DEFAULT_MANUSCRIPT = PROJECT_DIR / "docs" / "public_manuscript_template.md"
DEFAULT_OUT = CONTROLLED_ROOT / "hcc_ae_submission_package"
FIGURE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}


def set_table_geometry(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(10.5)
    styles["Heading 3"].font.bold = True


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_markdown_separator(line: str) -> bool:
    cells = parse_table_row(line)
    if not cells:
        return False
    return all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells)


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    rows = [row for row in rows if row and any(cell for cell in row)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            if row_idx == 0:
                shade_cell(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    set_table_geometry(table)
    doc.add_paragraph()


def add_markdown_to_docx(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        raw = lines[idx].rstrip()
        line = raw.strip()
        if not line:
            doc.add_paragraph()
            idx += 1
            continue
        if line.startswith("|"):
            table_rows: list[list[str]] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_line = lines[idx].strip()
                if not is_markdown_separator(table_line):
                    table_rows.append(parse_table_row(table_line))
                idx += 1
            add_table_from_rows(doc, table_rows)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)
        idx += 1


def iter_figure_files(figures_dir: Path) -> list[Path]:
    if not figures_dir.exists():
        raise FileNotFoundError(f"Figure directory was not found: {figures_dir}")
    return sorted(
        path
        for path in figures_dir.iterdir()
        if path.is_file() and path.suffix.lower() in FIGURE_SUFFIXES
    )


def add_figures(doc: Document, figures_dir: Path) -> list[Path]:
    figure_paths = iter_figure_files(figures_dir)
    if not figure_paths:
        return []
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for figure in figure_paths:
        doc.add_heading(figure.stem.replace("_", " "), level=2)
        doc.add_picture(str(figure), width=Inches(6.4))
    return figure_paths


def write_example_table_templates(out_dir: Path) -> list[Path]:
    tables = {
        "supplementary_table_1_cohort_template.csv": [
            "cohort_component",
            "n_slides",
            "n_cases",
            "source_category",
            "inclusion_note",
            "public_release_status",
        ],
        "supplementary_table_2_model_metrics_template.csv": [
            "model",
            "fold",
            "accuracy",
            "balanced_accuracy",
            "auc",
            "macro_f1",
            "notes",
        ],
        "supplementary_table_3_error_review_template.csv": [
            "anonymized_case",
            "error_pattern",
            "review_confirmed_label",
            "mimicry_score",
            "review_notes",
        ],
    }
    written: list[Path] = []
    for filename, header in tables.items():
        path = out_dir / filename
        with path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(header)
        written.append(path)
    return written


def write_manifest(
    out_dir: Path,
    manuscript: Path,
    docx_path: Path,
    figures_dir: Path | None,
    embedded_figures: list[Path],
    template_tables: list[Path],
) -> Path:
    manifest = out_dir / "build_manifest.txt"
    lines = [
        "HCC-AE public-safe package build",
        "",
        f"CONTROLLED_ROOT: {CONTROLLED_ROOT}",
        f"MANUSCRIPT_INPUT: {manuscript}",
        f"DOCX_OUTPUT: {docx_path}",
        f"FIGURES_DIR: {figures_dir if figures_dir else 'not provided'}",
        f"EMBEDDED_FIGURE_COUNT: {len(embedded_figures)}",
        f"TEMPLATE_TABLE_COUNT: {len(template_tables)}",
        "",
        "Release boundary:",
        "This builder does not include controlled data in the public repository.",
        "Generated outputs should remain outside public Git commits.",
        "Review author details, ethics, funding, source data, identifiers, and manuscript text before release.",
    ]
    if embedded_figures:
        lines.extend(["", "Embedded figures:"])
        lines.extend(str(path) for path in embedded_figures)
    if template_tables:
        lines.extend(["", "Example table templates:"])
        lines.extend(str(path) for path in template_tables)
    manifest.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return manifest


def build_docx(markdown: str, out_dir: Path, figures_dir: Path | None, embed_figures: bool) -> tuple[Path, list[Path]]:
    doc = Document()
    style_doc(doc)
    add_markdown_to_docx(doc, markdown)
    embedded_figures: list[Path] = []
    if embed_figures:
        if figures_dir is None:
            raise ValueError("--embed-figures requires --figures-dir")
        embedded_figures = add_figures(doc, figures_dir)
    docx_path = out_dir / "HCC_AE_submission.docx"
    doc.save(docx_path)
    return docx_path, embedded_figures


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a public-safe HCC-AE manuscript package from controlled local inputs."
    )
    parser.add_argument(
        "--manuscript",
        default=str(DEFAULT_MANUSCRIPT),
        help="Markdown manuscript input. Defaults to the public placeholder template.",
    )
    parser.add_argument(
        "--out",
        default=str(DEFAULT_OUT),
        help="Output directory. Prefer a private directory under HCC_AE_ROOT.",
    )
    parser.add_argument(
        "--figures-dir",
        default=None,
        help="Optional controlled local directory of figures to embed in the DOCX.",
    )
    parser.add_argument(
        "--embed-figures",
        action="store_true",
        help="Embed figures from --figures-dir in the DOCX. Figure files are not copied to Git.",
    )
    parser.add_argument(
        "--write-example-tables",
        action="store_true",
        help="Write header-only supplementary table templates to the output directory.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    manuscript = Path(args.manuscript).expanduser().resolve()
    out_dir = Path(args.out).expanduser().resolve()
    figures_dir = Path(args.figures_dir).expanduser().resolve() if args.figures_dir else None

    if not manuscript.exists():
        raise FileNotFoundError(f"Manuscript markdown was not found: {manuscript}")
    out_dir.mkdir(parents=True, exist_ok=True)

    markdown = manuscript.read_text(encoding="utf-8")
    docx_path, embedded_figures = build_docx(markdown, out_dir, figures_dir, args.embed_figures)
    template_tables = write_example_table_templates(out_dir) if args.write_example_tables else []
    manifest = write_manifest(out_dir, manuscript, docx_path, figures_dir, embedded_figures, template_tables)

    print(f"DOCX: {docx_path}")
    print(f"Manifest: {manifest}")
    print(f"Embedded figures: {len(embedded_figures)}")
    print(f"Example table templates: {len(template_tables)}")
    print("Public-safety note: generated outputs are controlled artifacts and should not be committed.")


if __name__ == "__main__":
    main()
