from __future__ import annotations

import csv
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(os.environ.get("FANLAB_RESEARCH2_ROOT", Path.cwd())).resolve()
PKG = ROOT / "final_manuscript_planning_20260605"
REV = PKG / "gut_microbes_revision"
MD_IN = REV / "gut_microbes_manuscript_expert_review.md"
DOCX_OUT = REV / "gut_microbes_manuscript_expert_review_with_figures.docx"
FIG_PREVIEWS = PKG / "compiled_figures" / "preview_pages"
FIG_MANIFEST = PKG / "compiled_figures" / "all_figures_one_pdf_manifest.csv"


def preferred_word_image(source: str, page: int) -> Path:
    src = Path(source)
    if src.suffix.lower() == ".png" and src.exists():
        return src
    png = src.with_suffix(".png")
    if png.exists():
        return png
    return FIG_PREVIEWS / f"page_{page:02d}.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_style_font(style, name: str, size_pt: float, color: str | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def _next_numbering_id(numbering, attr: str, tag: str) -> int:
    values = []
    for el in numbering.findall(qn(tag)):
        val = el.get(qn(attr))
        if val is not None:
            values.append(int(val))
    return (max(values) + 1) if values else 1


def new_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNumId", "w:abstractNum")
    num_id = _next_numbering_id(numbering, "w:numId", "w:num")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in [("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "left")]:
        child = OxmlElement(tag)
        child.set(qn("w:val"), val)
        lvl.append(child)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_paragraph_format(style, before=0, after=6, line=1.10, keep_next=False) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def configure_styles(doc: Document) -> None:
    # standard_business_brief preset: restrained, editable, manuscript-friendly.
    set_style_font(doc.styles["Normal"], "Calibri", 11, "000000")
    set_paragraph_format(doc.styles["Normal"], before=0, after=6, line=1.10)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        set_style_font(doc.styles[name], "Calibri", size, color, bold=True)
        set_paragraph_format(doc.styles[name], before=before, after=after, line=1.10, keep_next=True)

    set_style_font(doc.styles["Subtitle"], "Calibri", 10, "555555")
    set_paragraph_format(doc.styles["Subtitle"], before=0, after=8, line=1.10)

    set_style_font(doc.styles["List Bullet"], "Calibri", 11, "000000")
    set_paragraph_format(doc.styles["List Bullet"], before=0, after=4, line=1.10)
    set_style_font(doc.styles["List Number"], "Calibri", 11, "000000")
    set_paragraph_format(doc.styles["List Number"], before=0, after=4, line=1.10)


def set_page_geometry(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_text_runs(paragraph, text: str) -> None:
    # Supports the limited inline Markdown used in this manuscript: **bold** and `code`.
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_numbered_paragraph(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    apply_numbering(p, num_id)
    add_text_runs(p, text)


def add_markdown_paragraph(doc: Document, line: str) -> None:
    if line.startswith("# "):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string("0B2545")
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.05
        return

    if line.startswith("## "):
        doc.add_paragraph(line[3:], style="Heading 1")
        return

    if line.startswith("### "):
        doc.add_paragraph(line[4:], style="Heading 2")
        return

    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_text_runs(p, line[2:])
        return

    if line.startswith(">"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(line[1:].strip())
        run.italic = True
        run.font.color.rgb = RGBColor.from_string("1F4D78")
        return

    p = doc.add_paragraph()
    add_text_runs(p, line)


def add_manifest_table(doc: Document) -> None:
    rows = [
        ("Version", "Expert-review DOCX, editable text with figures placed near relevant Results sections"),
        ("Target journal", "Gut Microbes"),
        ("Administrative note", "Authors, ethics, consent, funding and final data availability are intentionally left for PI-team completion."),
        ("Evidence boundary", "Longitudinal association and candidate temporal links; no causal, mediation, regulation or strain-transmission claim."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.95)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v


def load_figure_manifest() -> dict[int, dict[str, str]]:
    figures: dict[int, dict[str, str]] = {}
    if not FIG_MANIFEST.exists():
        return figures
    with FIG_MANIFEST.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            figures[int(row["Page"])] = row
    return figures


FIGURE_CAPTIONS = {
    "Graphical abstract": "Graphical abstract. Longitudinal milk-gut-SCFA maturation model.",
    "Fig. 1": "Fig. 1. Study design and modular longitudinal multi-omics coverage.",
    "Fig. 2": "Fig. 2. Human milk bioactive component remodeling.",
    "Fig. 3": "Fig. 3. Infant gut microbiome maturation.",
    "Fig. 4": "Fig. 4. Infant fecal SCFA maturation.",
    "Fig. 5": "Fig. 5. Candidate lagged milk-to-SCFA links.",
    "Fig. S1": "Fig. S1. Maternal fecal metagenomic context.",
    "Fig. S2": "Fig. S2. Exploratory clinical modifier layer.",
}


def add_figure_block(doc: Document, figures: dict[int, dict[str, str]], page: int) -> None:
    row = figures.get(page)
    if not row:
        return
    figure = row["Figure"]
    image = preferred_word_image(row["Source"], page)
    if not image.exists():
        return

    p = doc.add_paragraph(style="Heading 2")
    p.add_run(figure).bold = True
    p.paragraph_format.keep_with_next = True

    with Image.open(image) as im:
        width_px, height_px = im.size
    ratio = width_px / height_px
    max_w = 6.1
    max_h = 7.0
    target_w = min(max_w, max_h * ratio)
    target_h = target_w / ratio
    if target_h > max_h:
        target_h = max_h
        target_w = max_h * ratio

    doc.add_picture(str(image), width=Inches(target_w), height=Inches(target_h))
    cap = doc.add_paragraph()
    cap.style = "Subtitle"
    cap.add_run(FIGURE_CAPTIONS.get(figure, figure))


INLINE_FIGURE_PLACEMENTS = {
    "Significance for Gut Microbes": [1],
    "Modular longitudinal coverage supports the analysis design": [2],
    "Human milk LTF and HMO features define the strongest exposure layer": [3],
    "Infant gut maturation is anchored by Bifidobacterium structure and SCFA maturation": [4, 5],
    "Early milk LTF/HMO features show prioritized candidate lagged links with later infant SCFAs": [6],
    "Maternal fecal profiles and clinical modifiers provide contextual support": [7, 8],
}


def heading_text(line: str) -> str | None:
    if line.startswith("## "):
        return line[3:]
    if line.startswith("### "):
        return line[4:]
    return None


def build_docx() -> None:
    doc = Document()
    set_page_geometry(doc)
    configure_styles(doc)

    # Footer page numbers are left to Word fields to keep the file editable.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Gut Microbes expert-review manuscript draft"
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = footer.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("777777")

    lines = MD_IN.read_text(encoding="utf-8").splitlines()
    figures = load_figure_manifest()
    title_done = False
    active_num_id: int | None = None
    current_heading: str | None = None
    placed_headings: set[str] = set()

    def flush_inline_figures() -> None:
        nonlocal current_heading
        if current_heading and current_heading in INLINE_FIGURE_PLACEMENTS and current_heading not in placed_headings:
            for page in INLINE_FIGURE_PLACEMENTS[current_heading]:
                add_figure_block(doc, figures, page)
            placed_headings.add(current_heading)

    for raw in lines:
        line = raw.rstrip()
        if not line:
            active_num_id = None
            continue
        this_heading = heading_text(line)
        if this_heading:
            flush_inline_figures()
            current_heading = this_heading
            active_num_id = None
        if not title_done and line.startswith("# "):
            add_markdown_paragraph(doc, line)
            add_manifest_table(doc)
            title_done = True
            active_num_id = None
            continue
        if line.startswith("|"):
            # The manuscript itself currently has no required Markdown tables; skip separators safely.
            active_num_id = None
            continue
        if re.match(r"^\d+\.\s+", line):
            if active_num_id is None:
                active_num_id = new_decimal_numbering(doc)
            add_numbered_paragraph(doc, re.sub(r"^\d+\.\s+", "", line), active_num_id)
            continue
        active_num_id = None
        add_markdown_paragraph(doc, line)

    flush_inline_figures()
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_docx()
